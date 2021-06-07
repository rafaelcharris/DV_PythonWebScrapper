'''
Requirements: Choose one news website - see article examples below for inspiration. 
Given a specific article URL from the website of your choice, 
return the title and content of the article to the user.

Parse out information such as the article title, updated date, and byline to return separately to the user.

The website I will scrape is la silla vacia

Let's add another feature. It writes a word document with the paper
'''

from docx import Document
import requests
from bs4 import BeautifulSoup
import re

#url = "https://lasillavacia.com/los-comunes-le-ofrecen-todo-petro-cambio-nada-concreto-81881"
url = input('Escriba la url de La Silla: ')

name_info = re.match("(https://lasillavacia.com/)(.*\D)", url)

name = name_info.group(2).replace("-", " ")[0:20]

r = requests.get(url)
soup = BeautifulSoup(r.text, 'html.parser')

# Get the artcle
article = soup.find_all("div", "body-node-historia")
autor_information = soup.find_all("div", "author author-top")

autor = autor_information[0].find("div", "editor").find("a").text

if autor is None:
    autor = "No disponible"

fecha_info = autor_information[0].find("div", "editor").find_all("span")[1].text
fecha = re.match(r' .\n(.*)', fecha_info).group(1)

if fecha is None:
    fecha = "No disponible"

# Create the word document that will store the info
document = Document()

document.add_heading(soup.title.text)
document.add_paragraph('Autor: ' + autor, style = 'List Bullet')
document.add_paragraph('Fecha: ' + fecha, style = 'List Bullet')
p = document.add_paragraph(article[0].text)

document.save(name + 'Scrapped.docx')
print("Documento Listo")
